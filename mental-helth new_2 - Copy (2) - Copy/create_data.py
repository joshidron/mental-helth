from docx import Document

def create_prajna_path_doc():
    doc = Document()
    
    # Title
    doc.add_heading('Prajñā Path', 0)
    doc.add_paragraph('Prajna Path')
    doc.add_paragraph('ભારતીય જ્ઞાન પરંપરાનો મર્મ દર્શાવે છે.')
    doc.add_paragraph('યુવા જીવનકૌશલ્ય માટે અત્યંત યોગ્ય માર્ગદર્શિકા.')

    # GAD
    doc.add_heading('ચિંતા અને ગભરાટની સમસ્યાઓ / Generalized Anxiety Disorder (GAD) / सामान्यीकृत चिंता विकार (GAD)', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / मुख्य लक्षण', level=2)
    doc.add_paragraph('અતિશય ચિંતા (Excessive worry) | अत्यधिक चिंता')
    doc.add_paragraph('ચંચળતા અથવા બેચેની (Restlessness) | बेचैनी')
    doc.add_paragraph('વધારે થાક લાગવો (Fatigue) | थकान')
    doc.add_paragraph('એકાગ્રતામાં મુશ્કેલી (Difficulty concentrating) | एकाग्रता में कठिनाई')
    doc.add_paragraph('ઊંઘમાં સમસ્યા (Sleep disturbance) | नींद की समस्या')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('સામાન્ય તણાવને જીએડી (GAD) માનવો (Thinking normal stress is GAD) | सामान्य तनाव को GAD मानना')
    doc.add_paragraph('શારીરિક સંવેદનાઓનું ખોટું અર્થઘટન (Misinterpreting body sensations) | शरीर की संवेदनाओं की गलत व्याख्या')
    doc.add_paragraph('ચિંતા એટલે વિકાર માનવો (Assuming worry = disorder) | चिंता को विकार मानना')
    doc.add_paragraph('એન્ઝાયટીનો ડર (Fear of having anxiety) | चिंता होने का डर')
    doc.add_paragraph('ઇન્ટરનેટથી સ્વ-નિદાન (Self-diagnosing from internet) | इंटरनेट से स्वयं निदान')
    
    # Depression
    doc.add_heading('ઉદાસીનતા / ડિપ્રેશન (Depression) / अवसाद (Depression)', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / मुख्य लक्षण', level=2)
    doc.add_paragraph('લગભગ રોજ અત્યંત ઉદાસી, ખાલીપણું અથવા નિરાશા (Persistent sad, empty, or hopeless mood) | लगातार उदासी, खालीपन या निराशा')
    doc.add_paragraph('રસ અથવા આનંદનો અભાવ (Loss of interest or pleasure) | गतिविधियों में रुचि या आनंद की कमी')
    doc.add_paragraph('ઊર્જાનો અભાવ અને થાક (Lack of energy and fatigue) | ऊर्जा की कमी और थकान')
    doc.add_paragraph('ભૂખ અને ઊંઘમાં ફેરફાર (Changes in appetite and sleep) | भूख और नींद में बदलाव')
    doc.add_paragraph('અપરાધભાવ અનુભવવો (Feeling worthless or guilty) | बेकार या दोषी महसूस करना')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('હળવી ઉદાસી (Mild sadness) | हल्की उदासी')
    doc.add_paragraph('આળસ (Laziness) | आलस')
    doc.add_paragraph('કામનો થાક (Work exhaustion) | काम की थकान')
    doc.add_paragraph('મૂડમાં બદલાવ (Mood swings) | मूड में बदलाव')
    doc.add_paragraph('સ્વ-વિશ્લેષણ (Self-analysis) | आत्म-विश्लेषण')
    

    # Stress
    doc.add_heading('તણાવ અને બર્નઆઉટ (Stress and Burnout) / तनाव और बर्नआउट', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / मुख्य लक्षण', level=2)
    doc.add_paragraph('અત્યંત શારીરિક અને માનસિક થાક (Overwhelming exhaustion) | अत्यधिक शारीरिक और मानसिक थकान')
    doc.add_paragraph('ચિડચિડાપણું અને મૂડ સ્વિંગ (Increased irritability) | चिड़चिड़ापन और मूड में बदलाव')
    doc.add_paragraph('એકાગ્રતામાં ઘટાડો (Reduced concentration) | एकाग्रता में कमी')
    doc.add_paragraph('રસ અને પ્રેરણા ગુમાવવી (Loss of interest and motivation) | रुचि और प्रेरणा की कमी')
    doc.add_paragraph('સામાજિક એકાંત (Social withdrawal) | सामाजिक दूरी')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('ક્ષણિક થાક (Temporary fatigue) | अस्थायी थकान')
    doc.add_paragraph('અસ્થાયી રુચિ ઘટવી (Temporary loss of interest) | अस्थायी अरुचि')
    doc.add_paragraph('નકારાત્મકતા (Negativity) | नकारात्मकता')

    # OCD
    doc.add_heading('બળજબરીના વિચારો અને વર્તણૂક (OCD) / जुनून-बाध्यकारी विकार (OCD)', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / मुख्य लक्षण', level=2)
    doc.add_paragraph('બળજબરીના વિચારો (Obsessions) | बाध्यकारी विचार')
    doc.add_paragraph('બળજબરીની હરકતો/વર્તણૂક (Compulsions) | बाध्यकारी व्यवहार')
    doc.add_paragraph('રોજબરોજની પ્રવૃત્તિઓમાં અડચણ (Interference with daily life) | दैनिक जीवन में बाधा')
    doc.add_paragraph('તાર્કિક ભાન હોવા છતાં નિયંત્રણ ન કરી શકવું (Loss of control despite logic) | तर्क के बावजूद नियंत्रण का अभाव')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('સુવ્યવસ્થા અને સ્વચ્છતાની પસંદગી (Preference for order and cleanliness) | व्यवस्था और स्वच्छता के लिए प्राथमिकता')
    doc.add_paragraph('વિશિષ્ટ ચિંતાઓ (Specific worries) | विशिष्ट चिंताएं')
    doc.add_paragraph('સામાન્ય ગેરવાજબી વિચારો (Normal irrational thoughts) | सामान्य तर्कहीन विचार')

    doc.add_heading('બળજબરીના વિચારો અને વર્તન ઘટાડવાના ઉપાયો (Ways to reduce OCD traits) / आवेशपूर्ण बाध्यकारी विकार (OCD) को कम करने के उपाय', level=2)
    
    doc.add_heading('1. સામાજિક પ્રવૃત્તિ (Social Activity) / सामाजिक गतिविधियाँ', level=2)
    doc.add_paragraph('એકલા ન રહો – પરિવાર અથવા મિત્રો સાથે વાત કરો (Do not stay alone; talk with family or friends) | अकेले न रहें; परिवार या मित्रों से बातचीत करें।')
    doc.add_paragraph('સમૂહ પ્રવૃત્તિઓમાં જોડાઓ: યોગ ક્લાસ, વૉકિંગ ગ્રુપ, ભજન/ધ્યાન (Join group activities: yoga, walking groups, meditation) | समूह गतिविधियों में शामिल हों: योग, वॉकिंग ग्रुप, ध्यान।')
    doc.add_paragraph('સેવા કાર્ય (વોલન્ટીયરિંગ) કરો – તે મનને અન્ય દિશામાં વાળે છે (Do volunteer work – it shifts focus outward) | स्वयंसेवी कार्य करें — इससे ध्यान बाहर की ओर जाता है।')
    doc.add_paragraph('વધુ પડતું સોશિયલ મીડિયા ટાળો – તે વિચારોમાં વધારો કરી શકે છે (Reduce social media usage; it may increase OCD thoughts) | सोशल मीडिया का कम उपयोग करें; यह OCD विचारों को बढ़ा सकता है।')

    doc.add_heading('2. કસરત (Exercise) / व्यायाम', level=2)
    doc.add_paragraph('દરરોજ 30–40 મિનિટ ચાલવું (Walk daily for 30–40 minutes) | प्रतिदिन 30–40 मिनट टहलें।')
    doc.add_paragraph('યોગ અને પ્રાણાયામ (Yoga and Pranayama) / योग और प्राणायाम:')
    doc.add_paragraph('- અનુલોમ-વિલોમ (Anulom–Vilom) / अनुलोम–विलोम')
    doc.add_paragraph('- ભ્રામરી (Bhramari) / भ्रामरी')
    doc.add_paragraph('- હળવું કપાલભાયી (Light Kapalbhati) / हल्का कपालभाति')
    doc.add_paragraph('હળવું સ્ટ્રેચિંગ – મન અને શરીર શાંત થાય છે (Gentle stretching relaxes the mind and body) | हल्की स्ट्रेचिंग मन और शरीर को शांत करती है।')

    doc.add_heading('3. આહાર શુદ્ધિ (Food Habits) / भोजन की आदतें', level=2)
    doc.add_paragraph('સંતુલિત આહાર લો (Eat a balanced diet) / संतुलित आहार लें:')
    doc.add_paragraph('- લીલી શાકભાજી, ફળ (Green vegetables and fruits) | हरी सब्जियाँ और फल')
    doc.add_paragraph('- દાળ, કઠોળ (Pulses and lentils) | दालें और अनाज')
    doc.add_paragraph('- બદામ, અખરોટ (Nuts like almonds and walnuts) | बादाम, अखरोट जैसे मेवे')
    doc.add_paragraph('ઓમેગા-3 માટે અળસી અને અખરોટ (Omega-3 foods like flaxseed, walnuts) | ओमेगा-3 युक्त भोजन (अलसी, अखरोट)')
    doc.add_paragraph('ચા-કોફીનું પ્રમાણ ઘટાડો (Reduce tea and coffee) | चाय और कॉफी का सेवन कम करें।')
    doc.add_paragraph('જંક ફૂડ અને વધુ પડતું તેલ કે મીઠું ટાળો (Avoid junk and processed food) | जंक और प्रोसेस्ड फूड से बचें।')
    doc.add_paragraph('પૂરતું પાણી પીવો (Drink enough water) | पर्याप्त मात्रा में पानी पिएँ।')

    doc.add_heading('4. દિનચર્યા અને માનસિક તકનીકો (Routine & Mind Techniques) / दिनचर्या और मन की तकनीकें', level=2)
    doc.add_paragraph('નિયમિત દિનચર્યા રાખો (Maintain a fixed daily routine) | नियमित दैनिक दिनचर्या बनाए रखें।')
    doc.add_paragraph('વિચારો સામે લડશો નહીં (Do not fight the thoughts) | विचारों से लड़ें नहीं।')
    doc.add_paragraph('સ્વયંને કહો: “આ માત્ર એક વિચાર છે, હકીકત નથી” (Tell yourself: “These are just thoughts, not reality”) | स्वयं से कहें: “ये सिर्फ विचार हैं, वास्तविकता नहीं।”')
    doc.add_paragraph('ધ્યાન અન્ય પ્રવૃત્તિ તરફ વાળો (Shift attention to another activity) | ध्यान को किसी अन्य गतिविधि में लगाएँ।')
    doc.add_paragraph('રોજ 5-10 મિનિટ તમારા વિચારો લખો (Journal for 5–10 minutes daily) | प्रतिदिन 5–10 मिनट डायरी लिखें।')
    doc.add_paragraph('10 મિનિટ ધ્યાન અથવા માઇન્ડફુલનેસ કરો (Practice mindfulness or meditation for 10 minutes) | 10 मिनट ध्यान (माइंडफुलनेस/मेडिटेशन) का अभ्यास करें।')

    # PTSD
    doc.add_heading('આઘાત બાદની પ્રતિક્રિયા (PTSD) / पोस्ट-ट्रॉमैटिक स्ट्रेस (PTSD)', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / मुख्य लक्षण', level=2)
    doc.add_paragraph('ઘટનાનું વારંવાર સ્મરણ કે અનુભવ થવો (Re-experiencing the event) | घटना का पुनः अनुभव')
    doc.add_paragraph('ટાળવું (Avoidance) | बचाव (सेहत के लिए टालना)')
    doc.add_paragraph('નકારાત્મક વિચારો (Negative thoughts) | नकारात्मक विचार')
    doc.add_paragraph('સજાગતામાં ફેરફારો (Changes in arousal) | सतर्कता में बदलाव')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('તાત્કાલિક શોક (Immediate grief) | तत्काल शोक')
    doc.add_paragraph('સામાન્ય ચિંતા (Normal anxiety) | सामान्य चिंता')
    doc.add_paragraph('ક્યારેક થતી ખરાબ યાદો (Occasional bad memory) | कभी-कभार होने वाली बुरी यादें')

    doc.add_heading('આઘાત પછીના તણાવને ઘટાડવાના ઉપાયો (Ways to Reduce Post-Traumatic Stress) / पोस्ट-ट्रॉमैटिक स्ट्रेस डिसऑर्डर (PTSD) को कम करने के उपाय', level=2)
    
    doc.add_heading('1. સામાજિક પ્રવૃત્તિઓ (Social Activities) / सामाजिक गतिविधियाँ', level=2)
    doc.add_paragraph('પરિવાર અને નજીકના મિત્રો સાથે ખુલ્લી વાતચીત કરવાથી માનસિક ભાર અને ભય ઓછો થાય છે. (Talking openly with family and close friends helps reduce emotional burden and fear.) | परिवार और करीबी दोस्तों से खुलकर बात करना भावनात्मक बोझ और डर को कम करता है।')
    doc.add_paragraph('સમૂહ પ્રવૃત્તિઓમાં ભાગ લેવાથી જોડાણ અને ભાવનાત્મક સુરક્ષાનો અનુભવ થાય છે. (Participating in group activities provides a sense of belonging and emotional safety.) | समूह गतिविधियों में भाग लेने से अपनापन और भावनात्मक सुरक्षा मिलती है।')
    doc.add_paragraph('સપોર્ટ ગ્રુપમાં જોડાવાથી વ્યક્તિ પોતાના અનુભવ શેર કરી શકે છે. (Joining support groups allows individuals to share experiences and feel understood.) | सपोर्ट ग्रुप्स में शामिल होने से अनुभव साझा करने और समझे जाने की भावना मिलती है।')
    doc.add_paragraph('સામાજિક સેવા અથવા વોલન્ટીયરિંગમાં જોડાવાથી ધ્યાન દુઃખદ યાદોથી દૂર જાય છે. (Engaging in community or volunteer work diverts attention from traumatic memories.) | सामुदायिक या स्वयंसेवी कार्य दर्दनाक यादों से ध्यान हटाते हैं।')
    doc.add_paragraph('સ્વસ્થ સામાજિક સંપર્ક દ્વારા એકાંતની ભાવના ઘટે છે અને સ્થિરતા વધે છે. (Reducing isolation through healthy social interaction improves emotional stability.) | स्वस्थ सामाजिक संपर्क अकेलेपन को कम कर भावनात्मक स्थिरता बढ़ाते हैं।')

    doc.add_heading('2. કસરત (Exercise) / व्यायाम', level=2)
    doc.add_paragraph('નિયમિત શારીરિક કસરત શરીરમાં જમા થયેલ તણાવ દૂર કરવામાં મદદ કરે છે. (Regular physical exercise helps release stress and tension stored in the body.) | नियमित व्यायाम शरीर में जमा तनाव को कम करता है।')
    doc.add_paragraph('યોગ અને સ્ટ્રેચિંગ મનને શાંત કરે છે અને ભાવનાત્મક સંતુલન આપે છે. (Yoga and stretching exercises promote relaxation and emotional balance.) | योग और स्ट्रेचिंग से आराम और भावनात्मक संतुलन मिलता है।')
    doc.add_paragraph('ચાલવું અથવા દોડવું જેવી પ્રવૃત્તિઓ મૂડ સુધારે છે અને ચિંતા ઘટાડે છે. (Aerobic activities like walking or jogging improve mood and reduce anxiety.) | वॉकिंग या जॉगिंग जैसे एरोबिक व्यायाम मूड सुधारते हैं और चिंता घटाते हैं।')
    doc.add_paragraph('શ્વાસની કસરતો ગભરાટ પર નિયંત્રણ રાખવામાં મદદરૂપ થાય છે. (Breathing exercises help control panic reactions and calm the nervous system.) | श्वसन अभ्यास घबराहट को नियंत्रित कर तंत्रिका तंत्र को शांत करते हैं।')
    doc.add_paragraph('નિયમિત કસરત ઊંઘની ગુણવત્તા સુધારે છે, જે PTSDમાં ઘણીવાર બગડેલી હોય છે. (Consistent exercise improves sleep quality, which is often disturbed in PTSD.) | निरंतर व्यायाम नींद की गुणवत्ता सुधारता है, जो PTSD में प्रभावित होती है।')

    doc.add_heading('3. આહાર શુદ્ધિ (Food Habits) / भोजन की आदतें', level=2)
    doc.add_paragraph('સંતુલિત આહાર મગજના સ્વાસ્થ્ય અને ભાવનાત્મક નિયંત્રણને મદદ કરે છે. (Eating a balanced diet supports brain health and emotional regulation.) | संतुलित आहार मस्तिष्क स्वास्थ्य और भावनात्मक नियंत्रण में सहायक है।')
    doc.add_paragraph('વિટામિન્સ અને ખનિજોથી ભરપૂર ખોરાક તણાવ અને થાક ઘટાડે છે. (Foods rich in vitamins and minerals help reduce stress and fatigue.) | विटामिन और खनिज युक्त भोजन तनाव और थकान कम करते हैं।')
    doc.add_paragraph('ઓમેગા-3યુક્ત ખોરાક માનસિક સ્પષ્ટતા અને સ્થિરતા વધારે છે. (Omega-3 rich foods improve mental clarity and emotional stability.) | ओमेगा-3 युक्त भोजन मानसिक स्पष्टता और भावनात्मक स्थिरता बढ़ाते हैं।')
    doc.add_paragraph('કેફીન અને વધુ પડતી ખાંડ લેવાનું ટાળો. (Limiting caffeine and sugar helps prevent anxiety and restlessness.) | कैफीन और शक्कर सीमित करने से बेचैनी और चिंता कम होती है।')
    doc.add_paragraph('પૂરતું પાણી પીવાથી શરીર અને મન બંને સંતુલિત રહે છે. (Drinking enough water keeps the body and mind well balanced.) | पर्याप्त पानी पीने से शरीर और मन संतुलित रहते हैं।')

    doc.add_heading('4. દિનચર્યા (Routine) / दिनचर्या', level=2)
    doc.add_paragraph('નિયમિત દિનચર્યા રાખવાથી જીવનમાં ગોઠવણ અને નિયંત્રણની ભાવના મળે છે. (Maintaining a fixed daily routine provides structure and a sense of control.) | नियमित दिनचर्या संरचना और नियंत्रण की भावना देती है।')
    doc.add_paragraph('સમયસર ઊંઘ અને જાગવાની ટેવ મૂડ અને ઉર્જાને સ્થિર રાખે છે. (Regular sleep and wake times help stabilize mood and energy levels.) | सोने-जागने का निश्चित समय मूड और ऊर्जा को स्थिर करता है।')
    doc.add_paragraph('દૈનિક કાર્યોનું આયોજન કરવાથી માનસિક ભાર ઓછો થાય છે. (Planning daily activities reduces uncertainty and mental overload.) | दैनिक योजना अनिश्चितता और मानसिक दबाव कम करती है।')
    doc.add_paragraph('દિનચર્યામાં આરામ માટે સમય ફાળવવાથી તણાવ નિયંત્રિત રહે છે. (Including relaxation time in the routine helps manage stress effectively.) | दिनचर्या में विश्राम का समय शामिल करने से तनाव बेहतर तरीके से संभलता है।')
    doc.add_paragraph('શિસ્તબદ્ધ જીવનશૈલી લાંબા ગાળાની સુધારણા અને માનસિક મજબૂતીમાં મદદ કરે છે. (A disciplined routine supports long-term recovery and emotional strength.) | अनुशासित दिनचर्या दीर्घकालिक सुधार और भावनात्मक मजबूती को बढ़ावा देती है।')

    # Personality
    doc.add_heading('વ્યક્તિત્વ અને ભાવનાત્મક અસંતુલન (Personality) / व्यक्तित्व और भावनात्मक असंतुलन', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / मुख्य लक्षण', level=2)
    doc.add_paragraph('અસ્થિર સંબંધો (Unstable relationships) | अस्थिर संबंध')
    doc.add_paragraph('અસ્થિર આત્મ-છબી (Unstable self-image) | अस्थिर आत्म-छवि')
    doc.add_paragraph('આવેગશીલતા (Impulsivity) | आवेग (Impulsivity)')
    doc.add_paragraph('ભાવનાત્મક અસ્થિરતા (Emotional instability) | भावनात्मक अस्थिरता')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('કિશોરાવસ્થાની અસ્થિરતા (Adolescent instability) | किशोरावस्था की अस्थिरता')
    doc.add_paragraph('તણાવ પ્રત્યેની પ્રતિક્રિયા (Reaction to stress) | तनाव के प्रति प्रतिक्रिया')
    doc.add_paragraph('સામાજિક ચિંતા (Social anxiety) | सामाजिक चिंता')

    doc.add_heading('વ્યક્તિત્વ અને ભાવનાત્મક અસંતુલન ઘટાડવાના ઉપાયો / व्यक्तित्व और भावनात्मक असंतुलन को कम करने के उपाय', level=2)

    doc.add_heading('1. સામાજિક પ્રવૃત્તિઓ (Social Activities) / सामाजिक गतिविधियाँ', level=2)
    doc.add_paragraph('સ્વસ્થ સામાજિક સંપર્ક સ્થિરતા અને આત્મવિશ્વાસ વધારવામાં મદદ કરે છે. (Healthy social interaction helps develop emotional stability and self-confidence.) | स्वस्थ सामाजिक संपर्क भावनात्मक स्थिरता और आत्म-विश्वास बढ़ाते हैं।')
    doc.add_paragraph('વિશ્વસનીય લોકો સાથે લાગણીઓ વહેંચવાથી માનસિક દબાણ ઘટે છે. (Sharing feelings with trusted people reduces emotional pressure and confusion.) | भरोसेमंद लोगों से भावनाएँ साझा करने से मानसिक दबाव और भ्रम कम होता है।')
    doc.add_paragraph('સમૂહ પ્રવૃત્તિઓમાં ભાગ લેવાથી સામાજિક કૌશલ્ય સુધરે છે. (Participating in group activities improves interpersonal skills and personality balance.) | समूह गतिविधियाँ व्यक्तित्व संतुलन और सामाजिक कौशल सुधारती हैं।')
    doc.add_paragraph('સકારાત્મક સામાજિક સહારો મૂડ સ્વિંગ અને ભાવનાત્મક પ્રતિક્રિયાઓ નિયંત્રિત કરે છે. (Positive social support helps control mood swings and emotional reactions.) | सकारात्मक सामाजिक सहयोग मूड स्विंग्स और भावनात्मक प्रतिक्रियाओं को नियंत्रित करता है।')
    doc.add_paragraph('એકાંત ટાળવાથી ભાવનાત્મક વિકાસ અને માનસિક સુખાકારી વધે છે. (Avoiding isolation promotes emotional growth and psychological well-being.) | अकेलेपन से बचाव भावनात्मक विकास और मानसिक स्वास्थ्य को बढ़ावा देता है।')

    doc.add_heading('2. કસરત (Exercise) / व्यायाम', level=2)
    doc.add_paragraph('નિયમિત કસરત ભાવનાઓને નિયંત્રિત કરવામાં મદદ કરે છે. (Regular physical exercise helps regulate emotions and reduce stress.) | नियमित व्यायाम भावनाओं को नियंत्रित करने और तनाव घटाने में मदद करता है।')
    doc.add_paragraph('યોગ અને ધ્યાન આત્મ-નિયંત્રણ અને જાગૃતિ વધારે છે. (Yoga and meditation improve self-control and emotional awareness.) | योग और ध्यान आत्म-नियंत्रण और भावनात्मक जागरूकता बढ़ाते हैं।')
    doc.add_paragraph('શારીરિક પ્રવૃત્તિથી મન પ્રફુલ્લિત રહે છે. (Physical activity releases endorphins, which enhance mood and positivity.) | शारीरिक गतिविधि एंडोर्फिन रिलीज करती है, जिससे मूड और सकारात्मकता बढ़ती है।')
    doc.add_paragraph('કસરત ગુસ્સો, ચિંતા અને ભાવનાત્મક તણાવ ઘટાડે છે. (Exercise helps reduce anger, anxiety, and emotional tension.) | व्यायाम क्रोध, चिंता और भावनात्मक तनाव को कम करता है।')
    doc.add_paragraph('સક્રિય જીવનશૈલી સંતુલિત અને સ્વસ્થ વ્યક્તિત્વને સહારો આપે છે. (A physically active lifestyle supports a balanced and healthy personality.) | सक्रिय जीवनशैली संतुलित और स्वस्थ व्यक्तित्व को समर्थन देती है।')

    doc.add_heading('3. આહાર શુદ્ધિ (Food Habits) / भोजन की आदतें', level=2)
    doc.add_paragraph('સંતુલિત પોષણ મગજની કાર્યક્ષમતા અને નિયંત્રણને મજબૂત બનાવે છે. (Balanced nutrition supports brain function and emotional regulation.) | संतुलित पोषण मस्तिष्क कार्य और भावनात्मक नियंत्रण में सहायक है।')
    doc.add_paragraph('વિટામિન્સ અને ખનિજોથી ભરપૂર ખોરાક મૂડ અને ઉર્જા સ્થિર રાખે છે. (Foods rich in vitamins and minerals help stabilize mood and energy levels.) | विटामिन और खनिज युक्त भोजन मूड और ऊर्जा स्तर को स्थिर करता है।')
    doc.add_paragraph('ઓમેગા-3 અને પ્રોટીનયુક્ત ખોરાક ભાવનાત્મક સંતુલન સુધારે છે. (Omega-3 and protein-rich foods improve emotional balance and focus.) | ओमेगा-3 और प्रोटीन युक्त भोजन भावनात्मक संतुलन और एकाग्रता बढ़ाते हैं।')
    doc.add_paragraph('જંક ફૂડ અને કેફીન ઘટાડવાથી સ્થિરતા વધે છે. (Limiting junk food, caffeine, and sugar reduces emotional instability.) | जंक फूड, कैफीन और शक्कर कम करने से भावनात्मक अस्थिरता घटती है।')
    doc.add_paragraph('પૂરતું પાણી પીવાથી માનસિક સ્પષ્ટતા અને શાંતિ જળવાય છે. (Proper hydration helps maintain mental clarity and calmness.) | पर्याप्त पानी पीने से मानसिक स्पष्टता और शांति बनी रहती है।')

    doc.add_heading('4. દિનચર્યા (Routine) / दिनचर्या', level=2)
    doc.add_paragraph('ગોઠવાયેલી દિનચર્યા સુરક્ષા અને આત્મશિસ્ત આપે છે. (A structured daily routine provides emotional security and self-discipline.) | सुव्यवस्थित दिनचर्या भावनात्मक सुरक्षा और आत्म-अनुशासन देती है।')
    doc.add_paragraph('નિયમિત ઊંઘની આદત માનસિક સંતુલન જાળવવામાં મદદ કરે છે. (Regular sleep patterns help maintain emotional control and mental balance.) | नियमित नींद भावनात्मक नियंत्रण और मानसिक संतुलन बनाए रखती है।')
    doc.add_paragraph('કાર્યોનું આયોજન કરવાથી આવેગી વર્તન ઘટે છે. (Planning daily activities reduces impulsive behavior and emotional chaos.) | दैनिक योजना आवेगपूर्ण व्यवहार और भावनात्मक अव्यवस्था को कम करती है।')
    doc.add_paragraph('દિનચર્યામાં આરામનો સમય ઉમેરવાથી ભાવનાત્મક પુનઃસ્થાપન થાય છે. (Including relaxation time in the routine supports emotional recovery.) | दिनचर्या में विश्राम का समय शामिल करने से भावनात्मक पुनर्प्राप्ति में सहायक होता है।')
    doc.add_paragraph('સતત રૂટિન વ્યક્તિત્વ વિકાસ અને ભાવનાત્મક પરિપક્વતા વધારે છે. (Consistent routines strengthen personality development and emotional maturity.) | निरंतर दिनचर्या व्यक्तित्व विकास और भावनात्मक परिपक्वता को मजबूत करती है।')

    # ADHD
    doc.add_heading('એડીએચડી (ADHD) / एडीएचडी (ADHD)', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / मुख्य लक्षण', level=2)
    doc.add_paragraph('ધ્યાન કેન્દ્રિત કરવામાં મુશ્કેલી (Difficulty concentrating) | ध्यान केंद्रित करने में कठिनाई')
    doc.add_paragraph('ભૂલકણાપણું (Forgetfulness) | भूलने की बीमारी')
    doc.add_paragraph('અતિસક્રિયતા (Hyperactivity) | अतिसक्रियता')
    doc.add_paragraph('આવેગશીલતા (Impulsivity) | आवेग (Impulsivity)')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('વિચલિતતા (Distractibility) | व्याकुलता')
    doc.add_paragraph('ઊંઘનો અભાવ (Lack of sleep) | नींद की कमी')

    # Autism
    doc.add_heading('ઓટિઝમ સ્પેક્ટ્રમ (Autism Spectrum) / ऑटिज़्म स्पेक्ट्रम (Autism)', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / મુખ્ય લક્ષણો', level=2)
    doc.add_paragraph('સામાજિક સંકેતોને સમજવામાં મુશ્કેલી (Difficulty understanding social cues) | सामाजिक संकेतों को समझने में कठिनाई')
    doc.add_paragraph('સંવાદ સાધવામાં પડકાર (Challenges in communication) | संचार शुरू करने में चुनौतियां')
    doc.add_paragraph('નિયમોમાં લવચીકતાનો અભાવ (Lack of flexibility in rules) | नियमों में लचीलेपन की कमी')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('માત્ર શરમાળ સ્વભાવ (Just shy nature) | शर्मीला स्वभाव')
    doc.add_paragraph('સામાજિક ચિંતા (Social anxiety) | सामाजिक चिंता')

    # Eating Disorders
    doc.add_heading('ખોરાક સંબંધિત સમસ્યાઓ (Eating Disorders) / खाने के विकार (Eating Disorders)', level=1)
    
    doc.add_heading('મુખ્ય લક્ષણો (Main Symptoms) / મુખ્ય લક્ષણો', level=2)
    doc.add_paragraph('વજન વિશેનો વિકૃત વિચાર (Distorted idea about weight) | वजन के बारे में विकृत विचार')
    doc.add_paragraph('ખાવાની વર્તણૂકમાં અસાધારણતા (Abnormal eating behavior) | खाने के असामान्य व्यवहार')
    
    doc.add_heading('ભ્રામક માન્યતાઓ (Myths) / भ्रामक धारणाएं', level=2)
    doc.add_paragraph('વજન ઘટાડવાનું ઝનૂન (Obsession with weight loss) | वजन घटाने का जुनून')
    doc.add_paragraph('ભાવનાત્મક રીતે ખાવું (Emotional eating) | भावनात्मक रूप से खाना')

    # Save
    doc.save('uploads/Prajna_Path.docx')
    print("Prajna_Path.docx created successfully in uploads/")

if __name__ == "__main__":
    create_prajna_path_doc()
